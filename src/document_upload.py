"""
Document upload handling for dynamic file processing in the RAG application.

This module provides functions to:
- Validate uploaded files (size, type, format)
- Process uploaded PDFs and text files
- Create vector stores from user-uploaded documents
- Manage temporary file storage and cleanup
- Handle upload errors gracefully

Functions:
    validate_upload_file(): Check file size and type.
    process_uploaded_file(): Extract text and create embeddings.
    create_upload_vector_store(): Build FAISS index from uploads.
    cleanup_upload_temp_files(): Remove temporary files.
"""

import os
import shutil
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

import streamlit as st
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS

# Import configuration for upload handling
from config import (
    ENABLE_DOCUMENT_UPLOAD,
    MAX_UPLOAD_FILE_SIZE_MB,
    MAX_TOTAL_UPLOAD_SIZE_MB,
    ALLOWED_FILE_EXTENSIONS,
    UPLOAD_TEMP_DIRECTORY,
    AUTO_CLEANUP_UPLOADS,
    MAX_PAGES_PER_UPLOAD,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
)

# Import embeddings and logging utilities
from src.embeddings import create_huggingface_embeddings # creates the embeddings for the uploaded documents
from src.logging_config import get_logger


def validate_upload_file(file: Any) -> tuple[bool, str]:
    """
    Validate that an uploaded file meets size and type requirements.

    This function checks:
    1. File extension is in ALLOWED_FILE_EXTENSIONS
    2. Individual file size <= MAX_UPLOAD_FILE_SIZE_MB
    3. File is not empty

    Args:
        file (Any): Streamlit uploaded file object.
            Contains: name, size, type attributes

    Returns:
        tuple[bool, str]: (is_valid, error_message)
            - (True, ""): File is valid, proceed with processing
            - (False, "error_message"): File failed validation, don't process

    Raises:
        None. Returns error in tuple instead.
    """

    if not ENABLE_DOCUMENT_UPLOAD:
        return False, "Document upload feature is disabled."

    # Extract file extension (case-insensitive)
    file_extension = Path(file.name).suffix.lower()

    if file_extension not in ALLOWED_FILE_EXTENSIONS:
        error_msg = (
            f"File type '{file_extension}' not allowed. "
            f"Supported types: {', '.join(ALLOWED_FILE_EXTENSIONS)}"
        )
        return False, error_msg
    
    # Get file size in MB
    file_size_mb = file.size / (1024 * 1024)

    if file_size_mb > MAX_UPLOAD_FILE_SIZE_MB:
        error_msg = (
            f"File size {file_size_mb:.2f} MB exceeds limit "
            f"of {MAX_UPLOAD_FILE_SIZE_MB} MB"
        )
        return False, error_msg

    if file.size == 0:
        error_msg = "Uploaded file is empty. Please upload a valid document."
        return False, error_msg

    return True, ""


def validate_upload_total_size(uploaded_files: list[Any]) -> tuple[bool, str]:
    """
    Validate that total uploaded file size doesn't exceed session limit.

    This function prevents users from uploading multiple large files that
    collectively exceed resource limits.

    Args:
        uploaded_files (list[Any]): List of uploaded file objects.

    Returns:
        tuple[bool, str]: (is_valid, error_message)
            - (True, ""): Total size within limits
            - (False, "error_message"): Total size exceeds limit
    """

    # calculate total size in MB
    total_size_mb = sum(f.size for f in uploaded_files) / (1024 * 1024)

    # check against total size limit
    if total_size_mb > MAX_TOTAL_UPLOAD_SIZE_MB:
        error_msg = (
            f"Total upload size {total_size_mb:.2f} MB exceeds session limit "
            f"of {MAX_TOTAL_UPLOAD_SIZE_MB} MB"
        )
        return False, error_msg

    return True, ""


def process_uploaded_file(file: Any) -> list[Document]:
    """
    Process an uploaded file and extract documents with content.

    This function handles different file types:
    - PDF: Extracts text using PyPDFLoader
    - TXT: Reads plain text
    - DOCX: Extracts text from Word documents

    Args:
        file (Any): The uploaded file to process.

    Returns:
        list[Document]: List of LangChain Document objects with:
            - page_content: The extracted text
            - metadata: Includes 'source' (filename) and 'page' (page number)

    Raises:
        ValueError: If file processing fails
        ImportError: If required library for file type is missing

    Side Effects:
        - Creates temporary files
        - Logs processing steps
        - May raise exceptions on processing failure
    """

    logger = get_logger() # logger to document processing steps
    documents: list[Document] = [] # stores the extracted documents from the uploaded file

    file_extension = Path(file.name).suffix.lower()  # determines the file type for processing

    # pdf files are processed using PyPDFLoader, which extracts text from each page and creates a Document object for each page
    if file_extension == ".pdf":
        try:
            # create a temporary directory for PDF processing
            with TemporaryDirectory() as temp_dir:
                # Write uploaded file to temp location
                temp_path = Path(temp_dir) / file.name
                temp_path.write_bytes(file.getvalue())

                # Load PDF with page limit
                try:
                    loader = PyPDFLoader(str(temp_path))
                    documents = loader.load()

                    # Validate documents were extracted
                    if not documents:
                        error_msg = f"No pages could be extracted from PDF {file.name}. The file may be empty, corrupted, or use an unsupported format."
                        raise ValueError(error_msg)

                    # Filter out empty documents
                    documents = [doc for doc in documents if doc.page_content and doc.page_content.strip()]
                    
                    if not documents:
                        error_msg = f"PDF {file.name} has no extractable text content (possibly scanned without OCR)"
                        raise ValueError(error_msg)

                    # Limit pages to prevent excessive processing
                    if len(documents) > MAX_PAGES_PER_UPLOAD:
                        documents = documents[:MAX_PAGES_PER_UPLOAD]

                except ValueError:
                    # Re-raise ValueError from document extraction
                    raise
                except Exception as e:
                    error_msg = f"Failed to load PDF: {str(e)}"
                    raise ValueError(error_msg)

        except ValueError:
            # Re-raise ValueError (validation or extraction error)
            raise
        except Exception as e:
            error_msg = f"PDF processing failed: {str(e)}"
            raise ValueError(error_msg)

    # process text files by reading the content and creating a single Document object
    elif file_extension == ".txt":
        try:
            # Read text file
            text_content = file.getvalue().decode("utf-8")

            # Validate text content is not empty
            if not text_content or not text_content.strip():
                error_msg = f"Text file {file.name} is empty or contains only whitespace"
                raise ValueError(error_msg)

            # Create a Document object
            document = Document(
                page_content=text_content,
                metadata={"source": file.name, "page": 0},
            )
            documents = [document]

        except UnicodeDecodeError as e:
            error_msg = f"Text file encoding error: {str(e)}"
            raise ValueError(error_msg)
        except ValueError:
            # Re-raise ValueError (validation error)
            raise

    # Process word documents using python-docx, extracting text from paragraphs and creating a Document object
    elif file_extension == ".docx":
        try:
            from docx import Document as DocxDocument
            
            # Create temporary file for docx processing
            with TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir) / file.name
                temp_path.write_bytes(file.getvalue())

                # Extract text from Word document
                doc = DocxDocument(str(temp_path))
                text_content = "\n".join([para.text for para in doc.paragraphs])

                # Validate text content is not empty
                if not text_content or not text_content.strip():
                    error_msg = f"Word document {file.name} is empty or contains no extractable text"
                    raise ValueError(error_msg)

                # Create Document object
                document = Document(
                    page_content=text_content,
                    metadata={"source": file.name, "page": 0},
                )
                documents = [document]

        except ImportError:
            error_msg = "python-docx library not installed. Install with: pip install python-docx"
            logger.error(error_msg)
            raise ImportError(error_msg)
        except ValueError:
            # Re-raise ValueError (validation error)
            raise
        except Exception as e:
            error_msg = f"Word document processing failed: {str(e)}"
            raise ValueError(error_msg)

    return documents # Return the list of Document objects extracted from the uploaded file


def create_upload_vector_store(
    uploaded_files: list[Any],
) -> tuple[FAISS, int, int, list[str], float]:
    """
    Create a FAISS vector store from uploaded files.

    This function:
    1. Validates all uploaded files
    2. Processes each file and extracts documents
    3. Splits documents into chunks
    4. Creates vector embeddings
    5. Builds FAISS index
    6. Returns index statistics and total file size

    Args:
        uploaded_files (list[Any]): List of files uploaded by user.

    Returns:
        tuple[FAISS, int, int, list[str], float]: 
            - FAISS: The created vector store
            - int: Total documents processed
            - int: Total chunks created
            - list[str]: List of successfully processed filenames
            - float: Total file size in MB

    Raises:
        ValueError: If any uploaded file fails validation or processing
        Exception: If vector store creation fails

    Side Effects:
        - Processes files (may take time for large uploads)
        - Creates temporary files and cleans up after
        - Logs all processing steps
    """

    logger = get_logger() # Logger to document the process logs
    
    all_documents: list[Document] = []
    successful_files: list[str] = []
    total_file_size_mb = 0.0

    is_valid, error_msg = validate_upload_total_size(uploaded_files)  # Validate total upload size against session limit
    if not is_valid:
        raise ValueError(error_msg)

    # Process each uploaded file
    for file in uploaded_files:
        # Validate individual file
        is_valid, error_msg = validate_upload_file(file)
        if not is_valid:
            continue

        # Process the file
        try:
            documents = process_uploaded_file(file)
            all_documents.extend(documents)
            successful_files.append(file.name)
            
            # Calculate file size
            file_size_mb = file.size / (1024 * 1024)
            total_file_size_mb += file_size_mb

            logger.info(
                f"Successfully added {file.name}: {len(documents)} documents ({file_size_mb:.2f} MB)"
            )

        except Exception as e:
            # Continue with next file instead of failing entirely
            continue

    # Check if any documents were successfully processed
    if not all_documents:
        error_msg = "No documents could be extracted from uploaded files."
        raise ValueError(error_msg)

    # Filter out any documents with empty or minimal content
    all_documents = [
        doc for doc in all_documents 
        if doc.page_content and len(doc.page_content.strip()) > 0
    ]
    
    # If all documents are empty after filtering, raise an error
    if not all_documents:
        error_msg = "All extracted documents are empty. Please check your file content."
        raise ValueError(error_msg)

    # Split documents into chunks for embedding
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    split_documents = text_splitter.split_documents(all_documents) # Split documents into chunks

    try:
        vector_store = FAISS.from_documents(
            split_documents,
            create_huggingface_embeddings(),
        )

        return vector_store, len(all_documents), len(split_documents), successful_files, total_file_size_mb

    except Exception as e:
        error_msg = f"Failed to create vector store: {str(e)}"
        logger.error(error_msg)
        raise


def cleanup_upload_temp_files() -> None:
    """
    Clean up temporary files created during upload processing.

    This function removes the upload temporary directory if:
    - AUTO_CLEANUP_UPLOADS is True
    - The directory exists

    Useful for freeing disk space and removing sensitive data
    after uploads have been processed.

    Returns:
        None

    Side Effects:
        - May delete files from disk
        - Logs cleanup operations
        - Does not raise errors if cleanup fails

    Examples:
        >>> # Call at app startup or after processing uploads
        >>> cleanup_upload_temp_files()
    """

    logger = get_logger()

    if not AUTO_CLEANUP_UPLOADS:
        logger.debug("Automatic upload cleanup disabled in config")
        return

    if UPLOAD_TEMP_DIRECTORY.exists():
        try:
            shutil.rmtree(UPLOAD_TEMP_DIRECTORY)
            logger.info(f"Cleaned up temporary upload directory: {UPLOAD_TEMP_DIRECTORY}")

        except Exception as e:
            logger.warning(
                f"Failed to cleanup temporary upload directory: {str(e)}"
            )


def get_upload_session_size() -> float:
    """
    Get the total size of uploads in current session.

    Returns the total size of all uploaded files that have been processed
    in this Streamlit session.

    Returns:
        float: Total uploaded size in MB (defaults to 0 if no uploads).

    Note:
        - Retrieves stored value from st.session_state.upload_session_size
        - Returns 0 if no files have been uploaded
        - Useful for displaying upload progress to user
    """

    # Get Stored Upload Session Size
    # This value was calculated and stored during upload processing
    total_size_mb = st.session_state.get("upload_session_size", 0.0)

    return total_size_mb


def get_upload_remaining_quota() -> float:
    """
    Get remaining upload quota for the session.

    Returns the amount of upload space remaining based on
    MAX_TOTAL_UPLOAD_SIZE_MB minus already uploaded files.

    Returns:
        float: Remaining upload quota in MB.
    """

    # Calculate Remaining Quota
    current_size_mb = get_upload_session_size()  # Current storage occupied by uploaded files
    remaining_mb = MAX_TOTAL_UPLOAD_SIZE_MB - current_size_mb  # Remaining quota based on session limit

    return max(0, remaining_mb)  # Prevents return of negative values
