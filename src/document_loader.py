"""
Document loader for handling multiple file types in the RAG application.

This module provides functions to load and validate documents from
the documents directory. It handles PDF, TXT, and DOCX files, parsing
them and converting them into LangChain Document objects suitable for processing.

Functions:
    load_pdf_documents(): Load all PDF, TXT, and DOCX files from the documents directory.
"""

from pathlib import Path
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document

# Import configuration for document directory path
from config import DOCUMENT_DIRECTORY


def load_pdf_documents() -> list[Document]:
    """
    Load all documents from the configured documents directory.

    This function loads ALL supported file types (.pdf, .txt, .docx) from 
    the DOCUMENT_DIRECTORY. Each file type is processed appropriately:
    - PDF: Parsed page-by-page using PyPDFLoader
    - TXT: Read as plain text
    - DOCX: Extracted using python-docx library

    The loaded documents include:
    - page_content (str): The extracted text content
    - metadata (dict): Including 'source' (file name) and file type info

    Returns:
        list[Document]: A list of LangChain Document objects from all loaded files.
            Empty list if no supported files are found.

    Raises:
        ValueError: If no documents are loaded from the directory.

    Note:
        - The DOCUMENT_DIRECTORY is configured in config.py
        - Supported extensions: .pdf, .txt, .docx
        - Each PDF page is a separate Document
        - TXT and DOCX files are each one Document
        - Subdirectories are NOT recursively searched (only top-level files)

    Examples:
        >>> documents = load_pdf_documents()
        >>> print(f"Loaded {len(documents)} documents")
        >>> print(documents[0].metadata)  # Shows file info
    """

    all_documents: list[Document] = []

    if not DOCUMENT_DIRECTORY.exists():
        raise ValueError(
            f"Documents directory not found at: {DOCUMENT_DIRECTORY}\n"
            f"Please create the directory and add PDF, TXT, or DOCX files."
        )

    pdf_files = list(DOCUMENT_DIRECTORY.glob("*.pdf"))
    if pdf_files:
        for pdf_file in pdf_files:
            try:
                loader = PyPDFLoader(str(pdf_file))
                documents = loader.load()
                if documents:
                    all_documents.extend(documents)
            except Exception as e:
                print(f"Warning: Failed to load PDF {pdf_file.name}: {str(e)}")

    txt_files = list(DOCUMENT_DIRECTORY.glob("*.txt"))
    for txt_file in txt_files:
        try:
            with open(txt_file, "r", encoding="utf-8") as f:
                text_content = f.read()
            if text_content.strip():
                doc = Document(
                    page_content=text_content,
                    metadata={"source": txt_file.name, "file_type": "text"}
                )
                all_documents.append(doc)
        except Exception as e:
            print(f"Warning: Failed to load text file {txt_file.name}: {str(e)}")

    docx_files = list(DOCUMENT_DIRECTORY.glob("*.docx"))
    if docx_files:
        try:
            from docx import Document as DocxDocument
            for docx_file in docx_files:
                try:
                    doc_obj = DocxDocument(str(docx_file))
                    text_content = "\n".join([para.text for para in doc_obj.paragraphs])
                    if text_content.strip():
                        doc = Document(
                            page_content=text_content,
                            metadata={"source": docx_file.name, "file_type": "docx"}
                        )
                        all_documents.append(doc)
                except Exception as e:
                    print(f"Warning: Failed to load DOCX {docx_file.name}: {str(e)}")
        except ImportError:
            print("Warning: python-docx not installed. DOCX files will be skipped.")
            print("Install with: pip install python-docx")

    if not all_documents:
        raise ValueError(
            f"No documents were loaded from the documents folder at: {DOCUMENT_DIRECTORY}\n"
            f"Supported formats: .pdf, .txt, .docx\n"
            f"Please add at least one supported document file."
        )

    return all_documents
